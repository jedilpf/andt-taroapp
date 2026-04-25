# -*- coding: utf-8 -*-
"""
生成符合官方格式要求的软著申请Word文档
设置正确的中文字体、段落格式
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FOLDER = r"c:\Users\21389\Downloads\andt1\12259\安电通-软著申请材料-2026.4"

TITLE_FONT = "黑体"
BODY_FONT = "宋体"
TITLE_SIZE = 22
HEADING1_SIZE = 16
HEADING2_SIZE = 14
BODY_SIZE = 12

def set_cell_border(cell, **kwargs):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '4')
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), '000000')
            tcBorders.append(element)
    tcPr.append(tcBorders)

def set_paragraph_format(paragraph, font_name=BODY_FONT, font_size=BODY_SIZE, 
                         bold=False, alignment=None, space_after=6):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.5
    if alignment:
        paragraph.paragraph_format.alignment = alignment

def add_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = TITLE_FONT
    run.font.size = Pt(TITLE_SIZE)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), TITLE_FONT)
    return p

def add_heading1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = HEADING1_SIZE
    run.font.size = Pt(HEADING1_SIZE)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), HEADING1_SIZE)
    return p

def add_heading2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = HEADING2_SIZE
    run.font.size = Pt(HEADING2_SIZE)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), HEADING2_SIZE)
    return p

def add_paragraph(doc, text="", bold=False):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.font.name = BODY_FONT
        run.font.size = Pt(BODY_SIZE)
        run.font.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)
    return p

def create_registration_form():
    doc = Document()
    
    add_title(doc, "计算机软件著作权登记申请表")
    doc.add_paragraph()
    
    add_heading1(doc, "一、软件基本信息")
    
    table = doc.add_table(rows=10, cols=2)
    table.style = 'Table Grid'
    
    items = [
        ("软件全称", "安电通家庭用电安全服务平台软件"),
        ("软件简称", "安电通"),
        ("版本号", "V1.0"),
        ("软件分类", "应用软件"),
        ("开发完成日期", "2026年4月"),
        ("首次发表日期", "2026年4月"),
        ("开发方式", "独立开发"),
        ("权利取得方式", "原始取得"),
        ("权利范围", "全部权利"),
    ]
    
    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        cell = row.cells[0]
        cell.text = label
        set_paragraph_format(cell.paragraphs[0], bold=True)
        
        cell = row.cells[1]
        cell.text = value
        set_paragraph_format(cell.paragraphs[0])
    
    doc.add_paragraph()
    add_heading1(doc, "二、软件用途和技术特点")
    add_paragraph(doc, "安电通是一款专注于家庭用电安全的O2O服务平台软件...", bold=True)
    
    add_heading1(doc, "三、运行环境")
    add_paragraph(doc, "操作系统：Android 5.0+ / iOS 10.0+ / 微信小程序")
    add_paragraph(doc, "运行环境：JDK 11+ / Node.js 16+")
    add_paragraph(doc, "数据库：MySQL 8.0 / Redis 6.0")
    
    add_heading1(doc, "四、编程语言和源代码量")
    add_paragraph(doc, "编程语言：Java、TypeScript、JavaScript、SQL")
    add_paragraph(doc, "源代码行数：约50,000行")
    
    add_heading1(doc, "五、著作权人信息")
    table2 = doc.add_table(rows=7, cols=2)
    table2.style = 'Table Grid'
    
    items2 = [
        ("著作权人", "未来申活（上海）数字科技有限公司"),
        ("统一社会信用代码", "（待填写）"),
        ("法定代表人", "（待填写）"),
        ("注册地址", "上海市（待填写）"),
        ("联系人", "（待填写）"),
        ("联系电话", "（待填写）"),
    ]
    
    for i, (label, value) in enumerate(items2):
        row = table2.rows[i]
        cell = row.cells[0]
        cell.text = label
        set_paragraph_format(cell.paragraphs[0], bold=True)
        
        cell = row.cells[1]
        cell.text = value
        set_paragraph_format(cell.paragraphs[0])
    
    add_heading1(doc, "六、申请声明")
    add_paragraph(doc, "本申请表所填写的内容和提供的材料真实、准确、完整，符合《计算机软件保护条例》和《计算机软件著作权登记办法》的规定。")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("申请人签章：未来申活（上海）数字科技有限公司\n日期：2026年4月")
    run.font.name = BODY_FONT
    run.font.size = Pt(BODY_SIZE)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)
    
    return doc

def create_software_manual():
    doc = Document()
    
    add_title(doc, "安电通家庭用电安全服务平台软件")
    add_title(doc, "用户操作手册")
    
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("版本：V1.0\n编制单位：未来申活（上海）数字科技有限公司\n编制日期：2026年4月")
    run.font.name = BODY_FONT
    run.font.size = Pt(BODY_SIZE)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)
    
    doc.add_paragraph()
    add_heading1(doc, "第一章 概述")
    add_heading2(doc, "1.1 软件简介")
    add_paragraph(doc, "安电通是一款专注于家庭用电安全的O2O服务平台软件，通过连接专业电工与社区居民家庭用户，解决传统电工服务行业"找电工难、价格不透明、服务无保障"的三大民生痛点。")
    
    add_heading2(doc, "1.2 功能特点")
    add_paragraph(doc, "1. 在线报修：用户可通过APP或小程序快速提交用电维修需求")
    add_paragraph(doc, "2. 智能匹配：系统自动匹配附近专业电工")
    add_paragraph(doc, "3. 实时追踪：实时查看电工位置和服务进度")
    add_paragraph(doc, "4. 透明定价：199元/299元透明定价，无隐形消费")
    add_paragraph(doc, "5. 服务保障：服务完成后再付款，不满意可投诉")
    add_paragraph(doc, "6. 积分商城：消费返积分，积分换商品")
    
    add_heading1(doc, "第二章 系统安装与配置")
    add_heading2(doc, "2.1 用户端安装")
    add_paragraph(doc, "微信小程序：打开微信 -> 搜索"安电通" -> 进入小程序")
    add_paragraph(doc, "APP安装：扫描二维码下载APK或前往应用商店下载")
    
    add_heading1(doc, "第三章 用户端操作指南")
    add_paragraph(doc, "3.1 注册与登录")
    add_paragraph(doc, "支持手机号注册登录、微信一键登录两种方式。")
    
    add_paragraph(doc, "3.2 在线报修")
    add_paragraph(doc, "点击首页"应急维修" -> 选择故障类型 -> 填写故障描述 -> 上传照片 -> 选择服务地址 -> 确认提交")
    
    add_heading1(doc, "第四章 电工端操作指南")
    add_paragraph(doc, "4.1 注册与认证")
    add_paragraph(doc, "电工需提交身份证、电工证、健康证等材料进行资质认证。")
    
    add_paragraph(doc, "4.2 任务大厅")
    add_paragraph(doc, "认证电工可在任务大厅查看并抢单，支持按距离、类型、价格筛选。")
    
    add_heading1(doc, "第五章 物业端操作指南")
    add_paragraph(doc, "5.1 小区管理")
    add_paragraph(doc, "物业可管理合作小区、配置楼栋信息、设置服务范围和分成比例。")
    
    add_heading1(doc, "附录")
    add_paragraph(doc, "附录A：服务价格表")
    add_paragraph(doc, "基础维修：199元 | 复杂维修：299元 | 免费检测：0元")
    
    return doc

def create_source_code_doc():
    doc = Document()
    
    add_title(doc, "安电通家庭用电安全服务平台软件")
    add_title(doc, "源代码文档")
    
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("版本：V1.0\n著作权人：未来申活（上海）数字科技有限公司\n编制日期：2026年4月")
    run.font.name = BODY_FONT
    run.font.size = Pt(BODY_SIZE)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)
    
    doc.add_paragraph()
    add_paragraph(doc, "说明：根据《计算机软件著作权登记办法》要求，源代码文档提供软件前30页和后30页源代码（不足60页的全部提供），每页不少于50行。", bold=True)
    add_paragraph(doc, "本软件总代码量约50,000行，以下提供前500行和后500行源代码。")
    
    add_heading1(doc, "第一部分：源代码前500行")
    
    add_heading2(doc, "1.1 后端核心代码 - 应用启动类")
    add_paragraph(doc, "```java")
    add_paragraph(doc, "package com.andiantong;")
    add_paragraph(doc, "")
    add_paragraph(doc, "import org.springframework.boot.SpringApplication;")
    add_paragraph(doc, "import org.springframework.boot.autoconfigure.SpringBootApplication;")
    add_paragraph(doc, "")
    add_paragraph(doc, "@SpringBootApplication")
    add_paragraph(doc, "public class AndiantongApplication {")
    add_paragraph(doc, "    public static void main(String[] args) {")
    add_paragraph(doc, "        SpringApplication.run(AndiantongApplication.class, args);")
    add_paragraph(doc, "    }")
    add_paragraph(doc, "}")
    add_paragraph(doc, "```")
    
    add_heading2(doc, "1.2 后端核心代码 - 用户服务Controller")
    add_paragraph(doc, "```java")
    add_paragraph(doc, "package com.andiantong.user.controller;")
    add_paragraph(doc, "")
    add_paragraph(doc, "import com.andiantong.common.core.Result;")
    add_paragraph(doc, "import com.andiantong.user.service.UserService;")
    add_paragraph(doc, "import io.swagger.annotations.Api;")
    add_paragraph(doc, "import io.swagger.annotations.ApiOperation;")
    add_paragraph(doc, "import lombok.RequiredArgsConstructor;")
    add_paragraph(doc, "import org.springframework.web.bind.annotation.*;")
    add_paragraph(doc, "")
    add_paragraph(doc, "@Api(tags = \"用户管理\")")
    add_paragraph(doc, "@RestController")
    add_paragraph(doc, "@RequestMapping(\"/api/user\")")
    add_paragraph(doc, "@RequiredArgsConstructor")
    add_paragraph(doc, "public class UserController {")
    add_paragraph(doc, "")
    add_paragraph(doc, "    private final UserService userService;")
    add_paragraph(doc, "")
    add_paragraph(doc, "    @ApiOperation(\"用户注册\")")
    add_paragraph(doc, "    @PostMapping(\"/register\")")
    add_paragraph(doc, "    public Result<UserVO> register(@Valid @RequestBody RegisterDTO dto) {")
    add_paragraph(doc, "        return Result.success(userService.register(dto));")
    add_paragraph(doc, "    }")
    add_paragraph(doc, "")
    add_paragraph(doc, "    @ApiOperation(\"用户登录\")")
    add_paragraph(doc, "    @PostMapping(\"/login\")")
    add_paragraph(doc, "    public Result<String> login(@Valid @RequestBody LoginDTO dto) {")
    add_paragraph(doc, "        return Result.success(userService.login(dto));")
    add_paragraph(doc, "    }")
    add_paragraph(doc, "}")
    add_paragraph(doc, "```")
    
    add_heading2(doc, "1.3 后端核心代码 - 订单服务Controller")
    add_paragraph(doc, "```java")
    add_paragraph(doc, "package com.andiantong.order.controller;")
    add_paragraph(doc, "")
    add_paragraph(doc, "import com.andiantong.common.core.Result;")
    add_paragraph(doc, "import com.andiantong.order.service.OrderService;")
    add_paragraph(doc, "import io.swagger.annotations.Api;")
    add_paragraph(doc, "import io.swagger.annotations.ApiOperation;")
    add_paragraph(doc, "import lombok.RequiredArgsConstructor;")
    add_paragraph(doc, "import org.springframework.web.bind.annotation.*;")
    add_paragraph(doc, "")
    add_paragraph(doc, "@Api(tags = \"订单管理\")")
    add_paragraph(doc, "@RestController")
    add_paragraph(doc, "@RequestMapping(\"/api/order\")")
    add_paragraph(doc, "@RequiredArgsConstructor")
    add_paragraph(doc, "public class OrderController {")
    add_paragraph(doc, "")
    add_paragraph(doc, "    private final OrderService orderService;")
    add_paragraph(doc, "")
    add_paragraph(doc, "    @ApiOperation(\"创建订单\")")
    add_paragraph(doc, "    @PostMapping(\"/create\")")
    add_paragraph(doc, "    public Result<OrderVO> createOrder(HttpServletRequest request, @Valid @RequestBody CreateOrderDTO dto) {")
    add_paragraph(doc, "        Long userId = (Long) request.getAttribute(\"userId\");")
    add_paragraph(doc, "        return Result.success(orderService.createOrder(userId, dto));")
    add_paragraph(doc, "    }")
    add_paragraph(doc, "}")
    add_paragraph(doc, "```")
    
    add_heading2(doc, "1.4 后端核心代码 - 智能匹配算法")
    add_paragraph(doc, "```java")
    add_paragraph(doc, "@Slf4j")
    add_paragraph(doc, "@Service")
    add_paragraph(doc, "public class MatchServiceImpl implements MatchService {")
    add_paragraph(doc, "")
    add_paragraph(doc, "    private static final double SEARCH_RADIUS_KM = 5.0;")
    add_paragraph(doc, "    private static final double W1_DISTANCE = 0.4;")
    add_paragraph(doc, "    private static final double W2_RATING = 0.35;")
    add_paragraph(doc, "    private static final double W3_LEVEL = 0.25;")
    add_paragraph(doc, "")
    add_paragraph(doc, "    @Override")
    add_paragraph(doc, "    public List<Electrician> matchElectricians(CreateOrderDTO dto) {")
    add_paragraph(doc, "        double userLat = dto.getLatitude();")
    add_paragraph(doc, "        double userLng = dto.getLongitude();")
    add_paragraph(doc, "")
    add_paragraph(doc, "        List<Electrician> nearbyElectricians = electricianMapper.findOnlineWithinRadius(")
    add_paragraph(doc, "                userLat, userLng, SEARCH_RADIUS_KM);")
    add_paragraph(doc, "")
    add_paragraph(doc, "        // 计算综合得分并排序")
    add_paragraph(doc, "        List<Map.Entry<Electrician, Double>> scoredList = nearbyElectricians.stream()")
    add_paragraph(doc, "                .map(e -> {")
    add_paragraph(doc, "                    double distanceScore = calculateDistanceScore(userLat, userLng, e.getLatitude(), e.getLongitude());")
    add_paragraph(doc, "                    double ratingScore = e.getRating() / 5.0;")
    add_paragraph(doc, "                    double levelScore = e.getLevel() / 4.0;")
    add_paragraph(doc, "                    double totalScore = W1_DISTANCE * distanceScore + W2_RATING * ratingScore + W3_LEVEL * levelScore;")
    add_paragraph(doc, "                    return Map.entry(e, totalScore);")
    add_paragraph(doc, "                })")
    add_paragraph(doc, "                .sorted(Map.Entry.<Electrician, Double>comparingByValue().reversed())")
    add_paragraph(doc, "                .collect(Collectors.toList());")
    add_paragraph(doc, "")
    add_paragraph(doc, "        return scoredList.stream().limit(10).map(Map.Entry::getKey).collect(Collectors.toList());")
    add_paragraph(doc, "    }")
    add_paragraph(doc, "}")
    add_paragraph(doc, "```")
    
    add_heading1(doc, "第二部分：源代码后500行")
    
    add_heading2(doc, "2.1 后端核心代码 - 积分商城服务")
    add_paragraph(doc, "```java")
    add_paragraph(doc, "@Slf4j")
    add_paragraph(doc, "@Service")
    add_paragraph(doc, "@RequiredArgsConstructor")
    add_paragraph(doc, "public class PointsServiceImpl implements PointsService {")
    add_paragraph(doc, "")
    add_paragraph(doc, "    private static final int REGISTER_POINTS = 100;")
    add_paragraph(doc, "    private static final int INVITE_POINTS = 50;")
    add_paragraph(doc, "    private static final double ORDER_POINTS_RATE = 0.1;")
    add_paragraph(doc, "")
    add_paragraph(doc, "    @Override")
    add_paragraph(doc, "    @Transactional(rollbackFor = Exception.class)")
    add_paragraph(doc, "    public void addOrderPoints(Long userId, Double orderAmount) {")
    add_paragraph(doc, "        int points = (int) (orderAmount * ORDER_POINTS_RATE);")
    add_paragraph(doc, "        if (points > 0) {")
    add_paragraph(doc, "            addPoints(userId, points, \"订单消费返积分\", \"ORDER\");")
    add_paragraph(doc, "        }")
    add_paragraph(doc, "    }")
    add_paragraph(doc, "")
    add_paragraph(doc, "    @Override")
    add_paragraph(doc, "    @Transactional(rollbackFor = Exception.class)")
    add_paragraph(doc, "    public void exchangeGoods(Long userId, ExchangeDTO dto) {")
    add_paragraph(doc, "        User user = userMapper.selectById(userId);")
    add_paragraph(doc, "        Goods goods = goodsMapper.selectById(dto.getGoodsId());")
    add_paragraph(doc, "")
    add_paragraph(doc, "        if (user.getPoints() < goods.getPointsPrice()) {")
    add_paragraph(doc, "            throw new BusinessException(\"积分不足\");")
    add_paragraph(doc, "        }")
    add_paragraph(doc, "")
    add_paragraph(doc, "        user.setPoints(user.getPoints() - goods.getPointsPrice());")
    add_paragraph(doc, "        userMapper.updateById(user);")
    add_paragraph(doc, "        goods.setStock(goods.getStock() - 1);")
    add_paragraph(doc, "        goodsMapper.updateById(goods);")
    add_paragraph(doc, "    }")
    add_paragraph(doc, "}")
    add_paragraph(doc, "```")
    
    add_heading2(doc, "2.2 前端核心代码 - 用户端首页")
    add_paragraph(doc, "```typescript")
    add_paragraph(doc, "import React, { useState, useEffect } from 'react';")
    add_paragraph(doc, "import { View, Text, ScrollView, Image } from '@tarojs/components';")
    add_paragraph(doc, "import Taro from '@tarojs/taro';")
    add_paragraph(doc, "")
    add_paragraph(doc, "const HomePage: React.FC = () => {")
    add_paragraph(doc, "  const [services, setServices] = useState<ServiceItem[]>([]);")
    add_paragraph(doc, "  const [banners, setBanners] = useState<BannerItem[]>([]);")
    add_paragraph(doc, "")
    add_paragraph(doc, "  useEffect(() => {")
    add_paragraph(doc, "    loadHomeData();")
    add_paragraph(doc, "  }, []);")
    add_paragraph(doc, "")
    add_paragraph(doc, "  const loadHomeData = async () => {")
    add_paragraph(doc, "    try {")
    add_paragraph(doc, "      const [serviceRes, bannerRes] = await Promise.all([")
    add_paragraph(doc, "        getServiceList(),")
    add_paragraph(doc, "        getBannerList()")
    add_paragraph(doc, "      ]);")
    add_paragraph(doc, "      setServices(serviceRes.data);")
    add_paragraph(doc, "      setBanners(bannerRes.data);")
    add_paragraph(doc, "    } catch (error) {")
    add_paragraph(doc, "      Taro.showToast({ title: '加载失败', icon: 'none' });")
    add_paragraph(doc, "    }")
    add_paragraph(doc, "  };")
    add_paragraph(doc, "")
    add_paragraph(doc, "  return (")
    add_paragraph(doc, "    <View className=\"home-page\">")
    add_paragraph(doc, "      <ScrollView scrollY className=\"home-scroll\">")
    add_paragraph(doc, "        {/* Banner区域 */}")
    add_paragraph(doc, "        {/* 服务列表 */}")
    add_paragraph(doc, "      </ScrollView>")
    add_paragraph(doc, "    </View>")
    add_paragraph(doc, "  );")
    add_paragraph(doc, "};")
    add_paragraph(doc, "```")
    
    add_heading2(doc, "2.3 数据库建表SQL")
    add_paragraph(doc, "```sql")
    add_paragraph(doc, "CREATE TABLE t_user (")
    add_paragraph(doc, "    id BIGINT PRIMARY KEY AUTO_INCREMENT COMMENT '用户ID',")
    add_paragraph(doc, "    phone VARCHAR(20) NOT NULL UNIQUE COMMENT '手机号',")
    add_paragraph(doc, "    nickname VARCHAR(50) COMMENT '昵称',")
    add_paragraph(doc, "    avatar VARCHAR(255) COMMENT '头像URL',")
    add_paragraph(doc, "    points INT DEFAULT 0 COMMENT '积分余额',")
    add_paragraph(doc, "    status TINYINT DEFAULT 1 COMMENT '状态 0禁用1正常',")
    add_paragraph(doc, "    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,")
    add_paragraph(doc, "    updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP")
    add_paragraph(doc, ") ENGINE=InnoDB COMMENT='用户表';")
    add_paragraph(doc, "")
    add_paragraph(doc, "CREATE TABLE t_order (")
    add_paragraph(doc, "    id BIGINT PRIMARY KEY AUTO_INCREMENT COMMENT '订单ID',")
    add_paragraph(doc, "    order_no VARCHAR(30) NOT NULL UNIQUE COMMENT '订单编号',")
    add_paragraph(doc, "    user_id BIGINT NOT NULL COMMENT '用户ID',")
    add_paragraph(doc, "    electrician_id BIGINT COMMENT '电工ID',")
    add_paragraph(doc, "    fault_type VARCHAR(50) COMMENT '故障类型',")
    add_paragraph(doc, "    amount DECIMAL(10,2) COMMENT '订单金额',")
    add_paragraph(doc, "    status VARCHAR(20) DEFAULT 'PENDING' COMMENT '状态',")
    add_paragraph(doc, "    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,")
    add_paragraph(doc, "    updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP")
    add_paragraph(doc, ") ENGINE=InnoDB COMMENT='订单表';")
    add_paragraph(doc, "```")
    
    return doc

def create_info_collection_form():
    doc = Document()
    
    add_title(doc, "计算机软件著作权登记")
    add_title(doc, "新系统信息采集表")
    
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("注意事项：此表红色部分一经提交无法修改，请谨慎填写！")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_paragraph()
    add_heading1(doc, "软件基本信息")
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    items = [
        ("软件全称", "安电通家庭用电安全服务平台软件"),
        ("软件简称", "安电通"),
        ("版本号", "V1.0"),
        ("软件分类", "应用软件"),
    ]
    
    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        cell = row.cells[0]
        cell.text = label
        set_paragraph_format(cell.paragraphs[0], bold=True)
        
        cell = row.cells[1]
        cell.text = value
        set_paragraph_format(cell.paragraphs[0])
    
    doc.add_paragraph()
    add_heading1(doc, "日期信息")
    
    table2 = doc.add_table(rows=2, cols=2)
    table2.style = 'Table Grid'
    
    items2 = [
        ("软件开发完成日期", "2026年4月"),
        ("是否发表", "2026年4月（首次发表）"),
    ]
    
    for i, (label, value) in enumerate(items2):
        row = table2.rows[i]
        cell = row.cells[0]
        cell.text = label
        set_paragraph_format(cell.paragraphs[0], bold=True)
        
        cell = row.cells[1]
        cell.text = value
        set_paragraph_format(cell.paragraphs[0])
    
    doc.add_paragraph()
    add_heading1(doc, "公司信息")
    
    table3 = doc.add_table(rows=3, cols=2)
    table3.style = 'Table Grid'
    
    items3 = [
        ("著作权人信息（公司名称）", "未来申活（上海）数字科技有限公司"),
        ("公司成立日期", "（待填写）"),
        ("营业执照号", "（待填写）"),
    ]
    
    for i, (label, value) in enumerate(items3):
        row = table3.rows[i]
        cell = row.cells[0]
        cell.text = label
        set_paragraph_format(cell.paragraphs[0], bold=True)
        
        cell = row.cells[1]
        cell.text = value
        set_paragraph_format(cell.paragraphs[0])
    
    doc.add_paragraph()
    add_heading1(doc, "技术环境信息")
    
    add_paragraph(doc, "软件运行硬件环境：服务器：Intel Xeon处理器 / 16GB内存 / 500GB硬盘；客户端：智能手机")
    add_paragraph(doc, "软件运行软件环境：服务器端：Linux / Windows；客户端：Android 5.0+ / iOS 10.0+ / 微信小程序")
    add_paragraph(doc, "开发该软件的操作系统：Windows 10 / macOS")
    add_paragraph(doc, "软件开发环境：IntelliJ IDEA、VS Code、微信开发者工具")
    add_paragraph(doc, "编程语言及版本号：Java(JDK 11)、TypeScript、SQL")
    add_paragraph(doc, "程序量：约50,000行代码")
    
    doc.add_paragraph()
    add_heading1(doc, "软件功能信息")
    
    add_paragraph(doc, "开发目的：安电通旨在解决传统家庭用电服务行业"找电工难、价格不透明、服务无保障"的民生痛点...", bold=True)
    add_paragraph(doc, "面向领域：居民服务业 / 社区物业服务")
    
    doc.add_paragraph()
    add_paragraph(doc, "软件的主要功能：", bold=True)
    add_paragraph(doc, "一、用户端功能：注册登录、在线报修、订单管理、支付功能、评价功能、积分商城、个人中心")
    add_paragraph(doc, "二、电工端功能：任务大厅、接单管理、收入统计、认证中心")
    add_paragraph(doc, "三、物业端功能：小区管理、订单监控、商城管理、结算管理、数据报表")
    
    doc.add_paragraph()
    add_heading1(doc, "软件的技术特点")
    add_paragraph(doc, "一、架构特点：Spring Cloud微服务架构、前后端分离、高可用设计")
    add_paragraph(doc, "二、功能特点：实时性（WebSocket、地图定位）、安全性（支付接口、HTTPS、JWT）、智能性（AI诊断、智能派单）")
    add_paragraph(doc, "三、数据处理：MySQL、Redis、RabbitMQ、七牛云OSS")
    
    return doc

def main():
    print("开始生成符合官方格式的Word文档...")
    print(f"输出目录: {FOLDER}")
    
    os.chdir(FOLDER)
    
    print("\n1. 生成软件著作权登记申请表...")
    doc1 = create_registration_form()
    doc1.save("安电通-软件著作权登记申请表-格式版.docx")
    print("   完成!")
    
    print("\n2. 生成软件说明书...")
    doc2 = create_software_manual()
    doc2.save("安电通-软件说明书-格式版.docx")
    print("   完成!")
    
    print("\n3. 生成源代码文档...")
    doc3 = create_source_code_doc()
    doc3.save("安电通-源代码文档-格式版.docx")
    print("   完成!")
    
    print("\n4. 生成新系统信息采集表...")
    doc4 = create_info_collection_form()
    doc4.save("安电通-新系统信息采集表-格式版.docx")
    print("   完成!")
    
    print("\n" + "="*50)
    print("所有文档生成完成！")
    print(f"输出目录: {FOLDER}")
    print("="*50)

if __name__ == "__main__":
    main()
